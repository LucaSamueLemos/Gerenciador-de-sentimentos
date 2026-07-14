"""
Diário de Emoções — aplicativo desktop feito com Flet.

Funcionalidades:
- Registro de emoções com descrição, intensidade, data e horário automáticos
- Histórico completo com exclusão de registros
- Gráfico de frequência das emoções
- Exportação dos registros para Word (.docx) e Excel (.xlsx)

Como rodar:
    pip install -r requirements.txt
    python main.py
"""

import sqlite3
from datetime import datetime, timedelta
from pathlib import Path

import flet as ft

# ---------------------------------------------------------------------------
# Configuração
# ---------------------------------------------------------------------------

APP_DIR = Path(__file__).parent
DB_PATH = APP_DIR / "emocoes.db"

EMOTIONS = [
    ("Feliz", "😊", "#FFD93D"),
    ("Triste", "😢", "#6C9BCF"),
    ("Ansioso", "😰", "#FF8C42"),
    ("Irritado", "😠", "#E74C3C"),
    ("Calmo", "😌", "#4FD1C5"),
    ("Animado", "🤩", "#F38181"),
    ("Cansado", "😴", "#A8A8B3"),
    ("Grato", "🙏", "#9B59B6"),
    ("Confuso", "😕", "#F1C40F"),
    ("Apaixonado", "🥰", "#FF6B9D"),
]
EMOTION_COLOR = {name: color for name, _, color in EMOTIONS}
EMOTION_EMOJI = {name: emoji for name, emoji, _ in EMOTIONS}

BG = "#0F0F1A"
SURFACE = "#1A1A2E"
SURFACE_2 = "#22223B"
PRIMARY = "#8C7CFF"
PRIMARY_2 = "#5EE7DF"
TEXT_MUTED = "#9A9AB3"


# ---------------------------------------------------------------------------
# Banco de dados
# ---------------------------------------------------------------------------

def init_db():
    con = sqlite3.connect(DB_PATH)
    con.execute(
        """
        CREATE TABLE IF NOT EXISTS emocoes (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            emocao TEXT NOT NULL,
            intensidade INTEGER NOT NULL,
            descricao TEXT,
            data TEXT NOT NULL,
            hora TEXT NOT NULL,
            timestamp TEXT NOT NULL
        )
        """
    )
    con.commit()
    con.close()


def insert_emotion(emocao: str, intensidade: int, descricao: str):
    now = datetime.now()
    con = sqlite3.connect(DB_PATH)
    con.execute(
        "INSERT INTO emocoes (emocao, intensidade, descricao, data, hora, timestamp) "
        "VALUES (?, ?, ?, ?, ?, ?)",
        (
            emocao,
            intensidade,
            descricao,
            now.strftime("%d/%m/%Y"),
            now.strftime("%H:%M"),
            # formato ordenável (SQLite compara strings lexicograficamente de forma correta)
            now.strftime("%Y-%m-%d %H:%M:%S"),
        ),
    )
    con.commit()
    con.close()


def _parse_timestamp(dt_str: str) -> datetime:
    # Compatibilidade: antes era ISO (datetime.now().isoformat()); agora é "YYYY-MM-DD HH:MM:SS".
    try:
        return datetime.fromisoformat(dt_str)
    except Exception:
        return datetime.strptime(dt_str, "%Y-%m-%d %H:%M:%S")



def fetch_all():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    rows = con.execute("SELECT * FROM emocoes ORDER BY timestamp DESC").fetchall()
    con.close()
    return [dict(r) for r in rows]


def fetch_all_filtered(start_dt: datetime | None, end_dt: datetime | None):
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row

    query = "SELECT * FROM emocoes WHERE 1=1 "
    params = []

    if start_dt is not None:
        query += " AND timestamp >= ? "
        params.append(start_dt.isoformat())

    if end_dt is not None:
        query += " AND timestamp <= ? "
        params.append(end_dt.isoformat())

    query += " ORDER BY timestamp DESC"

    rows = con.execute(query, params).fetchall()
    con.close()
    return [dict(r) for r in rows]


def fetch_counts_filtered(start_dt: datetime | None, end_dt: datetime | None):
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row

    query = "SELECT emocao, COUNT(*) as total FROM emocoes WHERE 1=1 "
    params = []

    if start_dt is not None:
        query += " AND timestamp >= ? "
        params.append(start_dt.isoformat())

    if end_dt is not None:
        query += " AND timestamp <= ? "
        params.append(end_dt.isoformat())

    query += " GROUP BY emocao ORDER BY total DESC"

    rows = con.execute(query, params).fetchall()
    con.close()
    return rows


def delete_emotion(record_id: int):
    con = sqlite3.connect(DB_PATH)
    con.execute("DELETE FROM emocoes WHERE id = ?", (record_id,))
    con.commit()
    con.close()


def fetch_counts():
    con = sqlite3.connect(DB_PATH)
    con.row_factory = sqlite3.Row
    rows = con.execute(
        "SELECT emocao, COUNT(*) as total FROM emocoes GROUP BY emocao ORDER BY total DESC"
    ).fetchall()
    con.close()
    return rows



# ---------------------------------------------------------------------------
# Exportação
# ---------------------------------------------------------------------------

def export_to_docx(path: str):
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    records = fetch_all()
    doc = Document()

    title = doc.add_heading("Diário de Emoções", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph(f"Exportado em {datetime.now().strftime('%d/%m/%Y às %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, text in enumerate(["Data", "Hora", "Emoção", "Intensidade", "Descrição"]):
        hdr[i].text = text
        hdr[i].paragraphs[0].runs[0].font.bold = True

    for r in records:
        row = table.add_row().cells
        row[0].text = r["data"]
        row[1].text = r["hora"]
        row[2].text = f'{EMOTION_EMOJI.get(r["emocao"], "")} {r["emocao"]}'
        row[3].text = f'{r["intensidade"]}/10'
        row[4].text = r["descricao"] or "-"

    doc.save(path)


def export_to_xlsx(path: str):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    records = fetch_all()
    wb = Workbook()
    ws = wb.active
    ws.title = "Emoções"

    headers = ["Data", "Hora", "Emoção", "Intensidade", "Descrição"]
    header_fill = PatternFill(start_color="8C7CFF", end_color="8C7CFF", fill_type="solid")
    header_font = Font(color="FFFFFF", bold=True, size=12)
    thin = Side(style="thin", color="DDDDDD")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = border

    for row_i, r in enumerate(records, start=2):
        values = [
            r["data"],
            r["hora"],
            f'{EMOTION_EMOJI.get(r["emocao"], "")} {r["emocao"]}',
            r["intensidade"],
            r["descricao"] or "-",
        ]
        for col_i, val in enumerate(values, start=1):
            cell = ws.cell(row=row_i, column=col_i, value=val)
            cell.border = border
            cell.alignment = Alignment(vertical="center", wrap_text=True)

    widths = [12, 10, 18, 12, 50]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    ws.freeze_panes = "A2"
    wb.save(path)


# ---------------------------------------------------------------------------
# Aplicativo
# ---------------------------------------------------------------------------

def main(page: ft.Page):
    page.title = "Diário de Emoções"
    page.bgcolor = BG
    page.window.width = 1100
    page.window.height = 760
    page.window.min_width = 820
    page.window.min_height = 600
    page.padding = 0
    page.fonts = {}
    page.theme = ft.Theme(font_family="Segoe UI")
    page.theme_mode = ft.ThemeMode.DARK

    init_db()

    # ---- estado ----
    selected_emotion = {"name": EMOTIONS[0][0]}
    snack = ft.SnackBar(content=ft.Text(""), bgcolor=SURFACE_2)
    page.overlay.append(snack)

    # filtro de tempo (afeta Histórico e Gráfico)
    time_filter = {"key": "all"}


    def notify(msg: str, color=PRIMARY_2):
        snack.content = ft.Text(msg, color="#FFFFFF")
        snack.bgcolor = SURFACE_2
        snack.open = True
        page.update()

    # ---- FilePickers para exportação ----
    def on_save_docx(e: ft.FilePickerResultEvent):
        if e.path:
            try:
                export_to_docx(e.path)
                notify("Exportado para Word com sucesso ✅")
            except Exception as ex:
                notify(f"Erro ao exportar: {ex}")

    def on_save_xlsx(e: ft.FilePickerResultEvent):
        if e.path:
            try:
                export_to_xlsx(e.path)
                notify("Exportado para Excel com sucesso ✅")
            except Exception as ex:
                notify(f"Erro ao exportar: {ex}")

    docx_picker = ft.FilePicker(on_result=on_save_docx)
    xlsx_picker = ft.FilePicker(on_result=on_save_xlsx)
    page.overlay.extend([docx_picker, xlsx_picker])

    def export_docx_click(e):
        docx_picker.save_file(
            dialog_title="Salvar registro em Word",
            file_name=f"emocoes_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
            allowed_extensions=["docx"],
        )

    def export_xlsx_click(e):
        xlsx_picker.save_file(
            dialog_title="Salvar registro em Excel",
            file_name=f"emocoes_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx",
            allowed_extensions=["xlsx"],
        )

    # =========================================================
    # VIEW 1 — REGISTRAR
    # =========================================================
    description_field = ft.TextField(
        label="Descreva o que você está sentindo",
        multiline=True,
        min_lines=4,
        max_lines=8,
        bgcolor=SURFACE_2,
        border_radius=14,
        border_color="transparent",
        focused_border_color=PRIMARY,
        color="#FFFFFF",
        label_style=ft.TextStyle(color=TEXT_MUTED),
    )

    intensity_slider = ft.Slider(
        min=1, max=10, divisions=9, value=5,
        active_color=PRIMARY, inactive_color=SURFACE_2,
        label="{value}",
    )
    intensity_text = ft.Text("5/10", size=14, color=TEXT_MUTED, weight=ft.FontWeight.W_600)

    def on_intensity_change(e):
        intensity_text.value = f"{int(intensity_slider.value)}/10"
        page.update()

    intensity_slider.on_change = on_intensity_change

    emotion_chips_row = ft.Row(wrap=True, spacing=10, run_spacing=10)

    def build_emotion_chips():
        emotion_chips_row.controls.clear()
        for name, emoji, color in EMOTIONS:
            is_selected = selected_emotion["name"] == name

            def make_click(n=name):
                def handler(e):
                    selected_emotion["name"] = n
                    build_emotion_chips()
                    page.update()
                return handler

            emotion_chips_row.controls.append(
                ft.Container(
                    content=ft.Row(
                        [ft.Text(emoji, size=18), ft.Text(name, size=13, weight=ft.FontWeight.W_600)],
                        spacing=6, tight=True,
                    ),
                    padding=ft.padding.symmetric(horizontal=16, vertical=10),
                    border_radius=30,
                    bgcolor=color if is_selected else SURFACE_2,
                    border=ft.border.all(2, color if is_selected else "transparent"),
                    on_click=make_click(),
                    ink=True,
                    animate=ft.Animation(200, "easeOut"),
                )
            )

    build_emotion_chips()

    def clear_form():
        description_field.value = ""
        intensity_slider.value = 5
        intensity_text.value = "5/10"
        selected_emotion["name"] = EMOTIONS[0][0]
        build_emotion_chips()
        page.update()

    def save_click(e):
        insert_emotion(
            selected_emotion["name"],
            int(intensity_slider.value),
            description_field.value.strip(),
        )
        notify(f"Registro salvo: {selected_emotion['name']} {EMOTION_EMOJI[selected_emotion['name']]}")
        clear_form()
        refresh_history()
        refresh_chart()

    clock_text = ft.Text(
        datetime.now().strftime("%A, %d de %B de %Y — %H:%M"),
        size=13, color=TEXT_MUTED,
    )

    register_view = ft.Container(
        padding=30,
        content=ft.Column(
            scroll=ft.ScrollMode.AUTO,
            spacing=20,
            controls=[
                ft.Text("Como você está se sentindo agora?", size=26, weight=ft.FontWeight.BOLD, color="#FFFFFF"),
                clock_text,
                ft.Container(height=10),
                ft.Text("Escolha uma emoção", size=14, color=TEXT_MUTED, weight=ft.FontWeight.W_600),
                emotion_chips_row,
                ft.Container(height=6),
                ft.Text("Intensidade", size=14, color=TEXT_MUTED, weight=ft.FontWeight.W_600),
                ft.Row([intensity_slider, intensity_text], spacing=10),
                description_field,
                ft.Container(height=6),
                ft.ElevatedButton(
                    content=ft.Row(
                        [ft.Icon(ft.Icons.SAVE_ROUNDED, size=20), ft.Text("Salvar registro", size=15, weight=ft.FontWeight.W_600)],
                        alignment=ft.MainAxisAlignment.CENTER, spacing=8,
                    ),
                    bgcolor=PRIMARY, color="#FFFFFF",
                    style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=14), padding=ft.padding.symmetric(vertical=18, horizontal=24)),
                    on_click=save_click,
                    width=260,
                ),
            ],
        ),
    )

    # =========================================================
    # VIEW 2 — HISTÓRICO
    # =========================================================
    history_list = ft.ListView(expand=True, spacing=12, padding=ft.padding.only(top=10))

    def make_delete_handler(record_id):
        def handler(e):
            delete_emotion(record_id)
            refresh_history()
            refresh_chart()
            notify("Registro excluído")
        return handler

    def get_filter_range():
        now = datetime.now()
        key = time_filter["key"]

        if key == "all":
            return None, None

        if key == "24h":
            return now.replace(microsecond=0) - timedelta(hours=24), now.replace(microsecond=0)
        if key == "7d":
            return now.replace(microsecond=0) - timedelta(days=7), now.replace(microsecond=0)
        if key == "30d":
            return now.replace(microsecond=0) - timedelta(days=30), now.replace(microsecond=0)

        if key == "month":
            start = now.replace(day=1, hour=0, minute=0, second=0, microsecond=0)
            return start, now.replace(microsecond=0)

        return None, None

    def refresh_history():
        history_list.controls.clear()
        start_dt, end_dt = get_filter_range()
        records = fetch_all_filtered(start_dt, end_dt)

        if not records:
            history_list.controls.append(
                ft.Container(
                    padding=40,
                    content=ft.Column(
                        [
                            ft.Icon(ft.Icons.SENTIMENT_SATISFIED_ALT, size=48, color=TEXT_MUTED),
                            ft.Text("Nenhum registro ainda", color=TEXT_MUTED, size=15),
                        ],
                        horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=10,
                    ),
                )
            )
        for r in records:
            color = EMOTION_COLOR.get(r["emocao"], PRIMARY)
            emoji = EMOTION_EMOJI.get(r["emocao"], "")
            history_list.controls.append(
                ft.Container(
                    bgcolor=SURFACE,
                    border_radius=16,
                    padding=18,
                    border=ft.border.only(left=ft.border.BorderSide(4, color)),
                    content=ft.Row(
                        alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                        vertical_alignment=ft.CrossAxisAlignment.START,
                        controls=[
                            ft.Row(
                                spacing=14,
                                controls=[
                                    ft.Container(
                                        content=ft.Text(emoji, size=26),
                                        bgcolor=SURFACE_2, border_radius=12, padding=10,
                                    ),
                                    ft.Column(
                                        spacing=3,
                                        controls=[
                                            ft.Row([
                                                ft.Text(r["emocao"], size=16, weight=ft.FontWeight.BOLD, color="#FFFFFF"),
                                                ft.Container(
                                                    content=ft.Text(f'{r["intensidade"]}/10', size=11, color="#FFFFFF", weight=ft.FontWeight.W_600),
                                                    bgcolor=color, border_radius=20,
                                                    padding=ft.padding.symmetric(horizontal=10, vertical=3),
                                                ),
                                            ], spacing=10),
                                            ft.Text(r["descricao"] or "Sem descrição", size=13, color=TEXT_MUTED, max_lines=3),
                                            ft.Text(f'{r["data"]} às {r["hora"]}', size=11, color=TEXT_MUTED),
                                        ],
                                    ),
                                ],
                            ),
                            ft.IconButton(
                                icon=ft.Icons.DELETE_OUTLINE_ROUNDED,
                                icon_color=TEXT_MUTED,
                                on_click=make_delete_handler(r["id"]),
                            ),
                        ],
                    ),
                )
            )
        page.update()

    def time_filter_chip(key: str, label: str):
        is_selected = time_filter["key"] == key

        def handler(e):
            time_filter["key"] = key
            # garante que o Flet atualize UI e dados
            refresh_history()
            refresh_chart()
            page.update()

        return ft.Container(
            content=ft.Text(
                label,
                size=12,
                color="#FFFFFF" if is_selected else TEXT_MUTED,
                weight=ft.FontWeight.W_600,
            ),
            padding=ft.padding.symmetric(horizontal=14, vertical=10),
            border_radius=30,
            bgcolor=PRIMARY if is_selected else SURFACE_2,
            border=ft.border.all(1, PRIMARY if is_selected else "transparent"),
            on_click=handler,
            ink=True,
        )


    # row de filtro reutilizável no Histórico e no Gráfico
    filter_row = ft.Row(
        wrap=True,
        spacing=10,
        run_spacing=10,
        controls=[
            time_filter_chip("all", "Tudo"),
            time_filter_chip("24h", "Últimas 24h"),
            time_filter_chip("7d", "Últimos 7 dias"),
            time_filter_chip("30d", "Últimos 30 dias"),
            time_filter_chip("month", "Este mês"),
        ],
    )

    history_view = ft.Container(
        padding=30,
        content=ft.Column(
            expand=True,
            controls=[
                ft.Row(
                    alignment=ft.MainAxisAlignment.SPACE_BETWEEN,
                    controls=[
                        ft.Text("Histórico", size=26, weight=ft.FontWeight.BOLD, color="#FFFFFF"),
                        ft.Row([
                            ft.OutlinedButton(
                                content=ft.Row([ft.Icon(ft.Icons.DESCRIPTION_OUTLINED, size=18), ft.Text("Exportar Word")], spacing=6, tight=True),
                                on_click=export_docx_click,
                                style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=12), side=ft.BorderSide(1, PRIMARY), color=PRIMARY),
                            ),
                            ft.OutlinedButton(
                                content=ft.Row([ft.Icon(ft.Icons.GRID_ON_ROUNDED, size=18), ft.Text("Exportar Excel")], spacing=6, tight=True),
                                on_click=export_xlsx_click,
                                style=ft.ButtonStyle(shape=ft.RoundedRectangleBorder(radius=12), side=ft.BorderSide(1, PRIMARY_2), color=PRIMARY_2),
                            ),
                        ], spacing=10),
                    ],
                ),
                ft.Container(height=6),
                filter_row,
                ft.Container(height=6),
                history_list,
            ],
        ),
    )


    # =========================================================
    # VIEW 3 — GRÁFICO
    # =========================================================
    chart_container = ft.Container(expand=True, alignment=ft.alignment.center)

    def refresh_chart():
        start_dt, end_dt = get_filter_range()
        counts = fetch_counts_filtered(start_dt, end_dt)

        if not counts:
            chart_container.content = ft.Column(
                [
                    ft.Icon(ft.Icons.BAR_CHART_ROUNDED, size=48, color=TEXT_MUTED),
                    ft.Text("Sem dados suficientes para exibir o gráfico", color=TEXT_MUTED),
                ],
                horizontal_alignment=ft.CrossAxisAlignment.CENTER, spacing=10,
                alignment=ft.MainAxisAlignment.CENTER,
            )
            page.update()
            return

        max_total = max(row[1] for row in counts)
        groups = []
        for i, row in enumerate(counts):
            emocao = row[0]
            total = row[1]

            color = EMOTION_COLOR.get(emocao, PRIMARY)

            groups.append(
                ft.BarChartGroup(
                    x=i,
                    bar_rods=[
                        ft.BarChartRod(
                            from_y=0, to_y=total, width=32,
                            color=color, tooltip=f"{emocao}: {total}",
                            border_radius=8,
                        )
                    ],
                )
            )

        chart = ft.BarChart(
            bar_groups=groups,
            border=ft.border.all(0, "transparent"),
            left_axis=ft.ChartAxis(labels_size=30, title=ft.Text("Registros"), title_size=20),
            bottom_axis=ft.ChartAxis(
                labels=[
                    ft.ChartAxisLabel(
                        value=i,
                        label=ft.Container(
                            ft.Text(f'{EMOTION_EMOJI.get(row["emocao"], "")}\n{row["emocao"]}', size=11, color=TEXT_MUTED, text_align=ft.TextAlign.CENTER),
                            padding=ft.padding.only(top=6),
                        ),
                    )
                    for i, row in enumerate(counts)
                ],
                labels_size=50,
            ),
            horizontal_grid_lines=ft.ChartGridLines(color=SURFACE_2, width=1),
            max_y=max_total + 1,
            interactive=True,
            expand=True,
        )

        total_registros = sum(r["total"] for r in counts)
        mais_frequente = counts[0]["emocao"]

        chart_container.content = ft.Column(
            expand=True,
            spacing=20,
            controls=[
                ft.Row(
                    spacing=16,
                    controls=[
                        stat_card("Total de registros", str(total_registros), PRIMARY),
                        stat_card("Emoção mais frequente", f'{EMOTION_EMOJI.get(mais_frequente, "")} {mais_frequente}', PRIMARY_2),
                        stat_card("Tipos de emoção registrados", str(len(counts)), "#F38181"),
                    ],
                ),
                ft.Container(
                    bgcolor=SURFACE, border_radius=18, padding=24, expand=True,
                    content=chart,
                ),
            ],
        )
        page.update()

    def stat_card(label, value, color):
        return ft.Container(
            expand=True,
            bgcolor=SURFACE, border_radius=16, padding=18,
            border=ft.border.only(top=ft.border.BorderSide(3, color)),
            content=ft.Column(
                spacing=4,
                controls=[
                    ft.Text(label, size=12, color=TEXT_MUTED),
                    ft.Text(value, size=20, weight=ft.FontWeight.BOLD, color="#FFFFFF"),
                ],
            ),
        )

    chart_view = ft.Container(
        padding=30,
        expand=True,
        content=ft.Column(
            expand=True,
            controls=[
                ft.Text("Gráfico de Emoções", size=26, weight=ft.FontWeight.BOLD, color="#FFFFFF"),
                ft.Container(height=10),
                filter_row,
                ft.Container(height=6),
                chart_container,
            ],
        ),
    )


    # =========================================================
    # NAVEGAÇÃO
    # =========================================================
    body = ft.Container(expand=True, content=register_view)

    def change_view(index):
        if index == 0:
            body.content = register_view
        elif index == 1:
            refresh_history()
            body.content = history_view
        else:
            refresh_chart()
            body.content = chart_view
        page.update()

    def on_nav_change(e):
        change_view(e.control.selected_index)

    nav_rail = ft.NavigationRail(
        selected_index=0,
        label_type=ft.NavigationRailLabelType.ALL,
        bgcolor=SURFACE,
        min_width=90,
        min_extended_width=180,
        on_change=on_nav_change,
        indicator_color=PRIMARY,
        destinations=[
            ft.NavigationRailDestination(icon=ft.Icons.ADD_CIRCLE_OUTLINE, selected_icon=ft.Icons.ADD_CIRCLE, label="Registrar"),
            ft.NavigationRailDestination(icon=ft.Icons.HISTORY_ROUNDED, selected_icon=ft.Icons.HISTORY_TOGGLE_OFF, label="Histórico"),
            ft.NavigationRailDestination(icon=ft.Icons.BAR_CHART_OUTLINED, selected_icon=ft.Icons.BAR_CHART_ROUNDED, label="Gráfico"),
        ],
    )

    header = ft.Container(
        padding=ft.padding.symmetric(horizontal=30, vertical=18),
        bgcolor=SURFACE,
        content=ft.Row(
            controls=[
                ft.Icon(ft.Icons.FAVORITE_ROUNDED, color=PRIMARY, size=26),
                ft.Text("Diário de Emoções", size=20, weight=ft.FontWeight.BOLD, color="#FFFFFF"),
            ],
            spacing=10,
        ),
    )

    page.add(
        ft.Column(
            expand=True, spacing=0,
            controls=[
                header,
                ft.Row(
                    expand=True, spacing=0,
                    controls=[
                        nav_rail,
                        ft.VerticalDivider(width=1, color=SURFACE_2),
                        body,
                    ],
                ),
            ],
        )
    )


if __name__ == "__main__":
    ft.app(target=main)
